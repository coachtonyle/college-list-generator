import os
import pandas as pd
import openai
from docx import Document
import shutil
import re
import time
from tqdm import tqdm
import logging
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

# Set up logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(levelname)s - %(message)s',
                    handlers=[logging.FileHandler("automation.log"),
                              logging.StreamHandler()])
logger = logging.getLogger(__name__)

# Constants
CSV_FILE = "hYLXTm.csv"
API_KEY = None  # Key will be set dynamically

def setup_openai(api_key):
    """Setup OpenAI client based on installed version"""
    try:
        # Try newer OpenAI library style
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        # Test if we can use the client this way
        client.models.list()
        logger.info("Using OpenAI API version 1.0.0+")
        return client, True
    except (ImportError, AttributeError):
        # Fall back to older style
        openai.api_key = api_key
        logger.info("Using OpenAI API version <1.0.0")
        return openai, False

def read_student_data(csv_file):
    """Read student data from CSV file"""
    try:
        df = pd.read_csv(csv_file)
        logger.info(f"Successfully loaded data for {len(df)} students")
        return df
    except Exception as e:
        logger.error(f"Error reading CSV file: {e}")
        return None

def set_cell_shading(cell, color):
    """Set shading for a table cell"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_border_to_table(table):
    """Add borders to all cells in a table"""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def create_section_heading(doc, text):
    """Create a section heading with underline formatting"""
    heading = doc.add_paragraph()
    run = heading.add_run(f"[{text}]")
    run.bold = True
    run.font.underline = True
    run.font.size = Pt(14)
    return heading

def generate_student_criteria_content(openai_client, student_info, is_new_api=False):
    """Generate personalized student criteria content"""
    try:
        # Extract student information
        first_name = student_info.get("Student's Name - First Name", "N/A")
        last_name = student_info.get("Student's Name - Last Name", "N/A")
        sat_score = student_info.get("SAT Score", "N/A")
        act_score = student_info.get("ACT Score", "N/A")
        majors = student_info.get("Please list the Majors you are considering", "N/A")
        career = student_info.get("Please list your Career Aspirations", "N/A")
        school_size = student_info.get("Size of the School: Do you prefer smaller classes with more personalized attention or larger classes with broader opportunities? Choose whether you prefer small classes (more interaction with professors, discussion-based learning) or large classes (more independence, lecture-style). Explain why. Do you prefer a small college, a medium-sized school, or a large university? Why?", "N/A")
        location = student_info.get("Location: Consider proximity to home, climate, urban/rural setting, and access to job markets or internships. Do you prefer an urban, suburban, or rural setting? Why? How close or far from home would you ideally like to be? (Close to home, A few hours away, Out of state, No preference)Are climate and weather important factors in your decision?  How important is campus safety and the surrounding area's security when considering a college?", "N/A")
        budget = student_info.get("Budget for Tuition and Fees: Compare costs and consider potential debt. What is your family's budget or comfort level with tuition costs? Are you aiming to minimize debt? Are you considering the cost of housing, food, transportation, and other living expenses when selecting a college?", "N/A")
        study_abroad = student_info.get("Study Abroad Options: Are study abroad programs or international education opportunities important to you?", "N/A")

        # Create prompt for student criteria section
        prompt = f"""
        Generate a detailed student criteria section for {first_name} {last_name}'s college recommendation document.
        
        Include ONLY the following information in bullet point format (each starting with a hyphen):
        - SAT Score: {sat_score}
        - ACT Score: {act_score} (if available)
        - Intended Majors: {majors}
        - Career Goals: {career}
        - Preferred School Type: (based on {school_size})
        - Location: (preferences based on {location})
        - Budget: (based on {budget})
        - Study Abroad: (based on {study_abroad})
        
        IMPORTANT: Format exactly as bullet points with ONLY the information above. Each bullet should start with a hyphen (-).
        Example format:
        - SAT Score: 1440
        - Intended Majors: Biology, Psychology
        - Career Goals: Medical professional
        
        Keep it concise and include ONLY these criteria points.
        """
        
        # Make API call based on version
        if is_new_api:
            response = openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert college counselor who creates personalized college plans."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=500
            )
            return response.choices[0].message.content.strip()
        else:
            response = openai_client.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert college counselor who creates personalized college plans."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=500
            )
            return response.choices[0].message['content'].strip()
    
    except Exception as e:
        logger.error(f"Error generating student criteria content: {e}")
        return None

def generate_overview_content(openai_client, student_info, is_new_api=False):
    """Generate personalized overview content"""
    try:
        first_name = student_info.get("Student's Name - First Name", "N/A")
        last_name = student_info.get("Student's Name - Last Name", "N/A")
        
        prompt = f"""
        Generate a personalized overview paragraph for {first_name} {last_name}'s college recommendation document.
        
        Create a paragraph that EXACTLY follows this structure:
        "This document presents a curated list of colleges that align with {first_name} {last_name}'s academic profile and the preferences shared. It serves as a starting point for their college admissions journey. Together, we will review this initial list to determine which institutions are the best fit and refine it into their final college list."
        
        Then add another paragraph exactly as follows:
        "Over the next couple of months, once we have finalized your list, it will guide our overall application strategy, including how we approach your personal statements and supplemental essays. As we move through the essay drafting and revision process, we will reference the specific prompts from the colleges you choose to apply to."
        
        And end with this exact paragraph:
        "Our goal is to finalize your college list by August 1st, ensuring we have a clear and focused path forward. We're excited to support you through this journey and look forward to a productive and thoughtful collaboration."
        
        IMPORTANT: Keep the exact phrasing provided and only personalize the student's name. Do not add any content beyond these three paragraphs.
        """
        
        # Make API call based on version
        if is_new_api:
            response = openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert college counselor who creates personalized college plans."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=500
            )
            return response.choices[0].message.content.strip()
        else:
            response = openai_client.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert college counselor who creates personalized college plans."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=500
            )
            return response.choices[0].message['content'].strip()
    
    except Exception as e:
        logger.error(f"Error generating overview content: {e}")
        return None

def generate_college_list(openai_client, student_info, is_new_api=False):
    """Generate personalized college list with reach, target, and safety schools"""
    try:
        first_name = student_info.get("Student's Name - First Name", "N/A")
        last_name = student_info.get("Student's Name - Last Name", "N/A")
        sat_score = student_info.get("SAT Score", "N/A")
        act_score = student_info.get("ACT Score", "N/A")
        majors = student_info.get("Please list the Majors you are considering", "N/A")
        career = student_info.get("Please list your Career Aspirations", "N/A")
        
        prompt = f"""
        Generate a comprehensive college list for {first_name} {last_name} with the following profile:
        - SAT: {sat_score}
        - ACT: {act_score}
        - Interested Majors: {majors}
        - Career Goals: {career}
        
        Create three categories of schools:
        1. Reach Schools (10-15 schools) - highly competitive institutions where admission is challenging
        2. Target Schools (6-10 schools) - schools where the student has a reasonable chance of admission
        3. Safety Schools (4-6 schools) - schools where admission is likely

        Format your response EXACTLY like this:
        
        Reach Schools
        - School Name 1
        - School Name 2
        - School Name 3
        (and so on)
        
        Target Schools
        - School Name 1
        - School Name 2
        - School Name 3
        (and so on)
        
        Safety Schools
        - School Name 1
        - School Name 2
        - School Name 3
        (and so on)
        
        Ensure each school offers programs in the student's areas of interest, and use full, formal university names (e.g., "University of California, Berkeley" not just "UC Berkeley").
        """
        
        # Make API call based on version
        if is_new_api:
            response = openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert college counselor who creates personalized college recommendations."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1000
            )
            return response.choices[0].message.content.strip()
        else:
            response = openai_client.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert college counselor who creates personalized college recommendations."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1000
            )
            return response.choices[0].message['content'].strip()
    
    except Exception as e:
        logger.error(f"Error generating college list: {e}")
        return None

def generate_reasons_for_selections(openai_client, student_info, college_list, is_new_api=False):
    """Generate reasons for school selections in each category"""
    try:
        first_name = student_info.get("Student's Name - First Name", "N/A")
        last_name = student_info.get("Student's Name - Last Name", "N/A")
        sat_score = student_info.get("SAT Score", "N/A")
        
        # Parse college list to extract schools by category
        reach_schools = []
        target_schools = []
        safety_schools = []
        
        current_category = None
        for line in college_list.split('\n'):
            line = line.strip()
            if "Reach Schools" in line:
                current_category = "reach"
            elif "Target Schools" in line:
                current_category = "target"
            elif "Safety Schools" in line:
                current_category = "safety"
            elif line.startswith('- ') and current_category:
                school = line[2:].strip()
                if current_category == "reach":
                    reach_schools.append(school)
                elif current_category == "target":
                    target_schools.append(school)
                elif current_category == "safety":
                    safety_schools.append(school)
        
        # Create prompt for reasons
        prompt = f"""
        Generate explanations for why each category of schools (Reach, Target, Safety) was selected for {first_name} {last_name} with an SAT score of {sat_score}.
        
        The specific schools in each category are:
        Reach Schools: {', '.join(reach_schools)}
        Target Schools: {', '.join(target_schools)}
        Safety Schools: {', '.join(safety_schools)}
        
        Write THREE separate paragraphs - one for REACH SCHOOLS, one for TARGET SCHOOLS, and one for SAFETY SCHOOLS - explaining:
        1. Why these schools are classified in this category (acceptance rates, SAT ranges relative to the student's score)
        2. Why they're still good matches for the student

        Format your response with these EXACT headers:
        
        **REACH SCHOOLS**
        
        [Paragraph explaining these schools]
        
        **TARGET SCHOOLS**
        
        [Paragraph explaining these schools]
        
        **SAFETY SCHOOLS**
        
        [Paragraph explaining these schools]
        
        Do NOT mention any student names other than {first_name} {last_name} in your response.
        Keep each paragraph focused ONLY on the general characteristics of that category.
        """
        
        # Make API call based on version
        if is_new_api:
            response = openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert college counselor who provides detailed justifications for college recommendations."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1000
            )
            return response.choices[0].message.content.strip()
        else:
            response = openai_client.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert college counselor who provides detailed justifications for college recommendations."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1000
            )
            return response.choices[0].message['content'].strip()
    
    except Exception as e:
        logger.error(f"Error generating reasons for selections: {e}")
        return None

def generate_detailed_college_info(openai_client, student_info, college_list, category, is_new_api=False):
    """Generate detailed information for colleges in a specific category"""
    try:
        first_name = student_info.get("Student's Name - First Name", "N/A")
        majors = student_info.get("Please list the Majors you are considering", "N/A")
        
        # Parse college list to extract schools for this category
        schools = []
        current_category = None
        for line in college_list.split('\n'):
            line = line.strip()
            if f"{category} Schools" in line:
                current_category = category.lower()
            elif line.startswith('- ') and current_category == category.lower():
                school = line[2:].strip()
                schools.append(school)
            elif current_category == category.lower() and any(cat in line for cat in ["Reach Schools", "Target Schools", "Safety Schools"]):
                break
        
        # Create prompt for detailed college info
        prompt = f"""
        Generate detailed information for the following {category.lower()} schools for {first_name}, who is interested in {majors}:
        
        {', '.join(schools[:5])}  # Limit to 5 schools to ensure detailed content within token limits
        
        For each college, provide the following information in this exact format:
        
        **College Name**
        College Level: {category}
        Location: [Full Address]
        College Type: [Private/Public]
        Ranking: [Ranking and source]
        College Overall Acceptance Rate: [Percentage]
        Specific Major Acceptance Rate: [Specific rate for the student's intended majors if available]
        Application Deadline Dates: [Early/Regular deadlines]
        SAT/ACT Range (middle 50%): [Score ranges]
        Recommended Majors: [3-4 relevant majors for this student]
        Application Fee: [Fee amount]
        Application Type: [Common App/Coalition/School specific]
        Tuition & Fees: [Cost breakdown]; Student-Faculty Ratio: [Ratio]
        Campus Life: [Brief description of campus environment]
        Pros and Cons:
        Pros: [3 pros]
        Cons: [2 cons]
        
        Provide accurate, up-to-date information for each school. Make sure all information is specific to {first_name} and their interests. Do not use any other student names in your response.
        """
        
        # Make API call based on version
        if is_new_api:
            response = openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert college counselor who provides detailed information about colleges."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=2000
            )
            return response.choices[0].message.content.strip()
        else:
            response = openai_client.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert college counselor who provides detailed information about colleges."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=2000
            )
            return response.choices[0].message['content'].strip()
    
    except Exception as e:
        logger.error(f"Error generating detailed college info: {e}")
        return None

def add_logo_to_document(doc, logo_path="company_logo.png"):
    """Add company logo to the header of every page"""
    try:
        # Access the sections of the document
        for section in doc.sections:
            # Access the header of the section
            header = section.header
            
            # Add a paragraph to the header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            
            # Add the logo to the header
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.5))  # Adjust width as needed
            
            # Center the logo
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        logger.info("Added company logo to document header")
        return True
    except Exception as e:
        logger.error(f"Error adding logo to document: {e}")
        return False


def create_new_document(student_first_name, student_last_name, logo_path="company_logo.png"):
    """Create a new blank document for the student with company logo in header"""
    try:
        doc = Document()
        output_filename = f"{student_first_name}_{student_last_name}_College_Recommendations.docx"
        
        # Check if file already exists and remove it
        if os.path.exists(output_filename):
            os.remove(output_filename)
        
        # Add the logo to the header if the logo file exists
        if os.path.exists(logo_path):
            add_logo_to_document(doc, logo_path)
        else:
            logger.warning(f"Logo file not found at {logo_path}. Header will not include logo.")
        
        logger.info(f"Creating new document: {output_filename}")
        return doc, output_filename
    except Exception as e:
        logger.error(f"Error creating new document: {e}")
        return None, None

def create_overview_section(doc, content):
    """Create the overview section in the document"""
    try:
        # Add section heading
        create_section_heading(doc, "OVERVIEW")
        
        # Add content paragraphs
        paragraphs = content.split('\n\n')
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                p = doc.add_paragraph()
                p.add_run(paragraph_text.strip())
        
        logger.info("Created overview section")
        return True
    
    except Exception as e:
        logger.error(f"Error creating overview section: {e}")
        return False

def create_student_criteria_section(doc, content):
    """Create the student criteria section in the document"""
    try:
        # Add section heading
        create_section_heading(doc, "STUDENT COLLEGE CRITERIA")
        
        # Add content as bullet points
        for line in content.split('\n'):
            line = line.strip()
            if line:
                p = doc.add_paragraph()
                
                if line.startswith('-'):
                    # Create bullet point with proper indentation
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.add_run("• " + line[1:].strip())
                else:
                    p.add_run(line)
        
        logger.info("Created student criteria section")
        return True
    
    except Exception as e:
        logger.error(f"Error creating student criteria section: {e}")
        return False

def create_college_list_table(doc, content, student_info=None):
    """Create the college list summary table in the document"""
    try:
        # Add section heading
        create_section_heading(doc, "COLLEGE LIST SUMMARY")
        
        # Create a table with 4 rows (header + 3 categories) and 2 columns
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add headers
        headers = table.rows[0].cells
        headers[0].text = "School Level"
        headers[1].text = "List of College"
        
        # Format headers
        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_shading(cell, "D3D3D3")
        
        # Add borders to the table
        add_border_to_table(table)
        
        # Parse the college list content
        current_category = None
        schools = []
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            if "Reach Schools" in line:
                if current_category and schools:
                    # Add previous category to table
                    row = table.add_row()
                    row.cells[0].text = current_category
                    row.cells[1].text = '\n'.join(schools)
                
                current_category = "Reach Schools"
                schools = []
            elif "Target Schools" in line:
                if current_category and schools:
                    # Add previous category to table
                    row = table.add_row()
                    row.cells[0].text = current_category
                    row.cells[1].text = '\n'.join(schools)
                
                current_category = "Target Schools"
                schools = []
            elif "Safety Schools" in line:
                if current_category and schools:
                    # Add previous category to table
                    row = table.add_row()
                    row.cells[0].text = current_category
                    row.cells[1].text = '\n'.join(schools)
                
                current_category = "Safety Schools"
                schools = []
            elif line.startswith('- ') and current_category:
                schools.append(line)
        
        # Add the last category
        if current_category and schools:
            row = table.add_row()
            row.cells[0].text = current_category
            row.cells[1].text = '\n'.join(schools)
        
        # Add a Student Picked Colleges row if needed
        if student_info and "Student Picked Colleges" in student_info:
            student_picks = student_info.get("Student Picked Colleges", [])
            if student_picks:
                row = table.add_row()
                row.cells[0].text = "Student Picked Colleges"
                row.cells[1].text = '\n'.join([f"- {college}" for college in student_picks])
        
        logger.info("Created college list summary table")
        return True
    
    except Exception as e:
        logger.error(f"Error creating college list table: {e}")
        return False

def create_reasons_section(doc, content):
    """Create the reasons for selections section in the document"""
    try:
        # Add section heading
        create_section_heading(doc, "REASONS FOR SELECTIONS")
        
        # Extract the categories from the content
        reach_content = ""
        target_content = ""
        safety_content = ""
        
        # Parse the content to extract the three sections
        current_section = None
        for line in content.split('\n'):
            if "**REACH SCHOOLS**" in line:
                current_section = "reach"
            elif "**TARGET SCHOOLS**" in line:
                current_section = "target"
            elif "**SAFETY SCHOOLS**" in line:
                current_section = "safety"
            elif current_section == "reach" and line.strip():
                reach_content += line + "\n"
            elif current_section == "target" and line.strip():
                target_content += line + "\n"
            elif current_section == "safety" and line.strip():
                safety_content += line + "\n"
        
        # Create a table to hold the reasons
        table = doc.add_table(rows=4, cols=2)  # 1 header row + 3 category rows
        
        # Set up headers
        headers = table.rows[0].cells
        headers[0].text = "School Category"
        headers[1].text = "Reasons for Selection"
        
        # Format headers
        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_shading(cell, "D3D3D3")
        
        # Add a border to the table
        add_border_to_table(table)
        
        # Add content for each category
        categories = [
            ("Reach Schools", reach_content.strip()),
            ("Target Schools", target_content.strip()),
            ("Safety Schools", safety_content.strip())
        ]
        
        for i, (category, content) in enumerate(categories):
            row = table.rows[i + 1]
            row.cells[0].text = category
            row.cells[1].text = content
            
            # Bold the category
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        logger.info("Created reasons for selections section")
        return True
    
    except Exception as e:
        logger.error(f"Error creating reasons section: {e}")
        return False

def create_detailed_college_info(doc, content_by_category):
    """Create the detailed college information section with all categories"""
    try:
        # Add section heading
        create_section_heading(doc, "COLLEGE DETAILED INFORMATION")
        
        # Dictionary for category order
        categories = {"Reach": 0, "Target": 1, "Safety": 2}
        
        # Process each category in order
        first_category = True
        for category in sorted(content_by_category.keys(), key=lambda k: categories.get(k, 99)):
            content = content_by_category[category]
            
            # Add a page break between categories (except before the first one)
            if not first_category:
                doc.add_page_break()
            else:
                first_category = False
            
            # Add a category heading
            heading = doc.add_heading(f"{category.upper()} SCHOOLS", level=1)
            for run in heading.runs:
                run.bold = True
            
            # Parse the content to extract college information
            colleges = []
            current_college = None
            college_details = ""
            
            for line in content.split('\n'):
                if line.startswith('**') and line.endswith('**'):
                    # If we have a previous college, add it to the list
                    if current_college and college_details:
                        colleges.append((current_college, college_details))
                        college_details = ""
                    
                    # This is a new college name
                    current_college = line.replace('*', '').strip()
                else:
                    # This is college details
                    college_details += line + "\n"
            
            # Add the last college
            if current_college and college_details:
                colleges.append((current_college, college_details))
            
            # Create a table for all colleges in this category
            if colleges:
                table = doc.add_table(rows=len(colleges) + 1, cols=2)
                
                # Add a border to the table
                add_border_to_table(table)
                
                # Add headers
                headers = table.rows[0].cells
                headers[0].text = "College Name"
                headers[1].text = "Details"
                
                # Format headers
                for cell in headers:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    set_cell_shading(cell, "D3D3D3")
                
                # Add content for each college
                for i, (college_name, details) in enumerate(colleges):
                    row = table.rows[i + 1]
                    row.cells[0].text = college_name
                    row.cells[1].text = details.strip()
                    
                    # Bold the college name
                    for paragraph in row.cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        
        logger.info("Created detailed college information section")
        return True
    
    except Exception as e:
        logger.error(f"Error creating detailed college info: {e}")
        return False

def extract_student_picked_colleges(student_info):
    """Extract student picked colleges from student information if available"""
    try:
        # Try to find any fields that might contain student college preferences
        potential_fields = [
            "Student Picked Colleges",
            "Student Selected Colleges",
            "Preferred Colleges",
            "Colleges of Interest"
        ]
        
        # Check for direct match in keys
        for field in potential_fields:
            if field in student_info and not pd.isna(student_info[field]):
                colleges = student_info[field]
                if isinstance(colleges, str):
                    # Split by common separators and clean up
                    college_list = [c.strip() for c in re.split(r'[,;\n]', colleges) if c.strip()]
                    return college_list
        
        # If we couldn't find a direct match, check for partial matches in keys
        for key in student_info:
            if any(field.lower() in key.lower() for field in potential_fields):
                if not pd.isna(student_info[key]):
                    colleges = student_info[key]
                    if isinstance(colleges, str):
                        college_list = [c.strip() for c in re.split(r'[,;\n]', colleges) if c.strip()]
                        return college_list
        
        # No student picked colleges found
        return []
    
    except Exception as e:
        logger.error(f"Error extracting student picked colleges: {e}")
        return []

# Update the create_college_list_table function to include student picked colleges
def create_college_list_table(doc, content, student_info=None):
    """Create the college list summary table in the document"""
    try:
        # Add section heading
        create_section_heading(doc, "COLLEGE LIST SUMMARY")
        
        # Determine number of rows needed (header + 3 categories + maybe student picks)
        student_picked_colleges = []
        if student_info:
            student_picked_colleges = extract_student_picked_colleges(student_info)
        
        initial_rows = 1  # header row
        has_student_picks = len(student_picked_colleges) > 0
        
        # Create a table
        table = doc.add_table(rows=initial_rows, cols=2)
        table.style = 'Table Grid'
        
        # Add headers
        headers = table.rows[0].cells
        headers[0].text = "School Level"
        headers[1].text = "List of College"
        
        # Format headers
        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_shading(cell, "D3D3D3")
        
        # Add borders to the table
        add_border_to_table(table)
        
        # Parse the college list content
        current_category = None
        schools = []
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            if "Reach Schools" in line:
                if current_category and schools:
                    # Add previous category to table
                    row = table.add_row()
                    row.cells[0].text = current_category
                    row.cells[1].text = '\n'.join(schools)
                
                current_category = "Reach Schools"
                schools = []
            elif "Target Schools" in line:
                if current_category and schools:
                    # Add previous category to table
                    row = table.add_row()
                    row.cells[0].text = current_category
                    row.cells[1].text = '\n'.join(schools)
                
                current_category = "Target Schools"
                schools = []
            elif "Safety Schools" in line:
                if current_category and schools:
                    # Add previous category to table
                    row = table.add_row()
                    row.cells[0].text = current_category
                    row.cells[1].text = '\n'.join(schools)
                
                current_category = "Safety Schools"
                schools = []
            elif line.startswith('- ') and current_category:
                schools.append(line)
        
        # Add the last category
        if current_category and schools:
            row = table.add_row()
            row.cells[0].text = current_category
            row.cells[1].text = '\n'.join(schools)
        
        # Add a Student Picked Colleges row if needed
        if has_student_picks:
            row = table.add_row()
            row.cells[0].text = "Student Picked Colleges"
            row.cells[1].text = '\n'.join([f"- {college}" for college in student_picked_colleges])
        
        logger.info("Created college list summary table")
        return True
    
    except Exception as e:
        logger.error(f"Error creating college list table: {e}")
        return False
    


def process_student(student_info, openai_client, is_new_api, logo_path="company_logo.png"):
    """Process a single student to create their personalized college recommendation document"""
    try:
        first_name = student_info.get("Student's Name - First Name", "")
        last_name = student_info.get("Student's Name - Last Name", "")
        
        if pd.isna(first_name) or pd.isna(last_name):
            logger.warning("Skipping student - missing name")
            return False
        
        logger.info(f"Processing student: {first_name} {last_name}")
        
        # Create a new document with company logo
        doc, output_file = create_new_document(first_name, last_name, logo_path)
        if not doc or not output_file:
            return False
        
        # Step 1: Generate and create the overview section
        logger.info("Generating overview content...")
        overview_content = generate_overview_content(openai_client, student_info, is_new_api)
        if overview_content:
            create_overview_section(doc, overview_content)
            # Save after each major section to prevent data loss
            doc.save(output_file)
        
        # Step 2: Generate and create the student criteria section
        logger.info("Generating student criteria content...")
        criteria_content = generate_student_criteria_content(openai_client, student_info, is_new_api)
        if criteria_content:
            create_student_criteria_section(doc, criteria_content)
            doc.save(output_file)
        
        # Step 3: Generate the college list
        logger.info("Generating college list...")
        college_list = generate_college_list(openai_client, student_info, is_new_api)
        if college_list:
            # Extract student picked colleges if available
            student_picked_colleges = extract_student_picked_colleges(student_info)
            
            # Create the college list table
            create_college_list_table(doc, college_list, {"Student Picked Colleges": student_picked_colleges})
            doc.save(output_file)
            
            # Step 4: Generate and create the reasons for selections
            logger.info("Generating reasons for selections...")
            reasons_content = generate_reasons_for_selections(openai_client, student_info, college_list, is_new_api)
            if reasons_content:
                create_reasons_section(doc, reasons_content)
                doc.save(output_file)
            
            # Step 5: Generate detailed college information for all categories
            content_by_category = {}
            
            for category in ["Reach", "Target", "Safety"]:
                logger.info(f"Generating detailed information for {category} schools...")
                detailed_info = generate_detailed_college_info(openai_client, student_info, college_list, category, is_new_api)
                if detailed_info:
                    content_by_category[category] = detailed_info
                
                # Sleep to avoid hitting API rate limits
                time.sleep(1)
            
            # Create the detailed college information section with all categories at once
            if content_by_category:
                create_detailed_college_info(doc, content_by_category)
                doc.save(output_file)
        
        # Final save
        doc.save(output_file)
        logger.info(f"Successfully created recommendation document for {first_name} {last_name}")
        return True
    
    except Exception as e:
        logger.error(f"Error processing student: {e}")
        return False

    


def main():
    """Main function to run the automation"""
    logger.info("College Recommendation Document Automation")
    logger.info("------------------------------------------")
    
    # Constants
    LOGO_PATH = "company_logo.png"  # Update this to your logo file path
    
    # Setup OpenAI client
    openai_client, is_new_api = setup_openai()
    
    # Read student data
    df = read_student_data(CSV_FILE)
    
    if df is not None:
        # Check if logo exists
        if not os.path.exists(LOGO_PATH):
            logger.warning(f"Logo file not found at {LOGO_PATH}. Documents will be created without logo.")
            user_input = input(f"Logo file not found at {LOGO_PATH}. Continue without logo? (y/n): ")
            if user_input.lower() != 'y':
                logger.info("Exiting program as requested.")
                return
        
        # Ask how many students to process
        try:
            user_input = input("How many students would you like to process? (Enter a number or 0 for all): ")
            num_students = int(user_input)
            if num_students > 0:
                df = df.head(num_students)
                logger.info(f"Processing {num_students} students")
            else:
                logger.info(f"Processing all {len(df)} students")
        except ValueError:
            logger.info(f"Invalid input, processing all {len(df)} students")
        
        # Process students with progress bar
        success_count = 0
        for index, row in tqdm(df.iterrows(), total=len(df), desc="Processing students"):
            # Process student with logo
            if process_student(row.to_dict(), openai_client, is_new_api, LOGO_PATH):
                success_count += 1
            
            # Sleep to avoid hitting API rate limits
            time.sleep(2)
        
        # Report results
        logger.info(f"Processing complete! Successfully processed {success_count} out of {len(df)} students.")
    else:
        logger.error("Failed to read student data. Please check your CSV file.")

if __name__ == "__main__":
    main()